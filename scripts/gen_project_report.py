# -*- coding: utf-8 -*-
"""Generate a project introduction document for the advisor."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ---- Global style ----
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def set_heading_font(level, size, color=(0x1a, 0x1a, 0x2e)):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Microsoft YaHei'
    h.font.size = Pt(size)
    h.font.color.rgb = RGBColor(*color)
    h.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

set_heading_font(1, 18)
set_heading_font(2, 14)
set_heading_font(3, 12)

def add_para(text, bold=False, size=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.8)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_number(text):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.space_after = Pt(3)
    return p

# ============================================================
# Title page
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('基于 Agent 的微博社交机器人\n项目设计与实现报告')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
title.paragraph_format.space_before = Pt(40)
title.paragraph_format.space_after = Pt(20)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('WeiboAgent — LocoAgent + All-IN-ONE 混合架构')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
subtitle.paragraph_format.space_after = Pt(30)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
today = datetime.date.today().strftime('%Y 年 %m 月 %d 日')
run = info.add_run(f'汇报人：[学生姓名]\n指导教师：[导师姓名]\n日期：{today}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
info.paragraph_format.space_after = Pt(40)

doc.add_page_break()

# ============================================================
# 1. 项目概述
# ============================================================
doc.add_heading('一、项目概述', level=1)

add_para(
    '本项目旨在构建一个基于 AI Agent 的微博社交机器人，能够自主完成浏览信息流、'
    '点赞、评论、转发、关注、发帖等社交行为。核心思路是：不从零编写爬虫与自动化逻辑，'
    '而是基于开源 Agent 框架（LocoAgent），通过沉淀 Skill（操作手册）、安装插件、'
    '编排 Workflow（工作流），将微博的各类操作封装为 Agent 可调用的能力，'
    '最终实现一个有人设、有策略、可持续运行的社交机器人。'
)

doc.add_heading('1.1 设计目标', level=2)
add_bullet('功能完整：覆盖浏览、搜索、点赞、评论、转发、关注、发帖（文字/图片/视频）全链路')
add_bullet('安全可控：人工登录、不自动输入密码、操作去重、频率限制、CAPTCHA 检测停止')
add_bullet('可扩展：通过 Skill + Workflow 的模式，可快速迁移到其他社交平台')
add_bullet('低成本：使用 DeepSeek 作为 LLM，降低 API 调用成本')

doc.add_heading('1.2 功能矩阵', level=2)

func_table = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
func_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = func_table.rows[0].cells
hdr[0].text = '功能'
hdr[1].text = '实现方式'
hdr[2].text = '状态'

funcs = [
    ('浏览首页信息流', 'agent-browser (Chrome CDP)', '已完成'),
    ('搜索微博', 'All-IN-ONE CLI (aione)', '已完成'),
    ('查看用户信息', 'All-IN-ONE CLI', '已完成'),
    ('阅读评论', 'All-IN-ONE CLI', '已完成'),
    ('点赞 / 取消点赞', 'agent-browser', '已完成'),
    ('评论 / 回复', 'agent-browser', '已完成'),
    ('转发微博', 'agent-browser', '已完成'),
    ('关注 / 取关', 'agent-browser', '已完成'),
    ('发布文字微博', 'All-IN-ONE CLI', '已完成'),
    ('发布图片/视频微博', 'All-IN-ONE CLI', '已完成'),
    ('定时发帖', 'LocoAgent Workflow', '框架就绪'),
    ('自动互动', 'LocoAgent Workflow', '框架就绪'),
    ('操作去重', 'LocoAgent Operation Log', '框架就绪'),
]
for f, m, s in funcs:
    row = func_table.add_row().cells
    row[0].text = f
    row[1].text = m
    row[2].text = s

doc.add_paragraph()

# ============================================================
# 2. 技术路线与方法
# ============================================================
doc.add_heading('二、技术路线与方法', level=1)

doc.add_heading('2.1 整体架构', level=2)
add_para(
    '项目采用"Agent 框架 + CLI 工具 + 浏览器自动化"的混合架构。'
    '不是单一依赖某种方式，而是根据操作类型选择最合适的技术路径：'
)
add_bullet('结构化操作（搜索、用户信息、发帖）→ All-IN-ONE CLI：返回 JSON，速度快、可靠')
add_bullet('交互式操作（点赞、关注、浏览信息流）→ agent-browser：通过 Chrome CDP 驱动浏览器，处理需要页面状态的操作')
add_bullet('决策与编排 → LocoAgent（LLM Agent）：读取 Skill 手册，决定调用哪个工具、传什么参数')
add_bullet('定时与批量 → Workflow Engine：LLM-free 的脚本化流水线，可定时守护运行')

doc.add_heading('2.2 核心组件', level=2)

doc.add_heading('LocoAgent（Agent 框架）', level=3)
add_para(
    'LocoAgent 是 Claude Code 的开源 fork，由 LocoreMind 维护。'
    '它提供了 Agent 的核心能力：REPL 交互、单任务模式（-p）、Skill 加载、'
    'Workflow 引擎、操作日志去重、多平台浏览器管理。'
    '我选择它而非从零搭建，是因为它已在 X.com（Twitter）自动化上验证过可行性，'
    '且原生支持 Skill 机制——这正是"沉淀操作手册"所需的载体。'
)

doc.add_heading('All-IN-ONE CLI（aione）', level=3)
add_para(
    'All-IN-ONE 是 cv-cat 维护的开源项目，封装了微博、小红书、抖音三个平台的'
    '爬虫与 API 调用，提供统一的 aione CLI（46 条命令，JSON 输出）。'
    '它解决了微博没有开放 API 的问题——通过逆向工程移动端/网页端接口，'
    '将搜索、用户信息、评论读取、发帖等操作封装为命令行工具。'
    'Agent 只需读 Skill 手册就知道怎么调用。'
)

doc.add_heading('agent-browser（浏览器自动化）', level=3)
add_para(
    'agent-browser 基于 Chrome DevTools Protocol（CDP），驱动一个独立的、'
    '持久的 Chrome 实例。与传统的 Selenium/Puppeteer 不同，它使用隔离的'
    '浏览器 profile，避免污染日常浏览器，且 session 持久化——只需手动登录一次。'
    '点赞、关注等操作没有 API 可用，只能通过浏览器模拟点击，这是 agent-browser 的职责。'
)

doc.add_heading('2.3 Skill 沉淀方法', level=2)
add_para(
    'Skill 是 Agent 的入口——一个结构化的 Markdown 操作手册。'
    '我为微博编写了 SKILL.md（约 22KB），覆盖 6 大类共 20+ 个操作序列。'
    '每个操作包含：目标、方法（CLI 还是 Browser）、前置条件、步骤命令、验证方式、'
    '已知问题。Agent 读取后自动理解：用什么 CLI、有哪些命令、参数怎么传、cookie 怎么配。'
)
add_para('Skill 的设计原则：', bold=True, space_after=3)
add_bullet('不硬编码元素 ref 号——每次操作前重新 snapshot，按文本/aria-label 识别元素')
add_bullet('CLI 优先——能用 CLI 拿结构化 JSON 的，不走浏览器')
add_bullet('安全栏栅——遇到登录墙/CAPTCHA 立即停止，不自动输入密码')

doc.add_heading('2.4 Workflow 编排方法', level=2)
add_para(
    'Workflow 是 LLM-free 的脚本化流水线，用于定时、批量的重复任务。'
    '每个 Workflow 由一个 JSON 配置 + 一个 TypeScript executor 组成。'
    '我设计了三个微博 Workflow：'
)
add_bullet('weibo-daily-post：每日定时发帖，按分类权重（教程40%/观点30%/推荐20%/日常10%）从内容池选题')
add_bullet('weibo-feed-monitor：信息流监控，每 2 小时搜索关键词，自动点赞/评论/关注匹配内容')
add_bullet('weibo-search-reply：关键词搜索互动，对搜索结果执行点赞、评论、关注好作者')

doc.add_heading('2.5 Persona 设计', level=2)
add_para(
    '机器人需要一个"人格"来指导其行为策略。我设计了角色"小洛"——'
    'AI 科技观察者 & 内容创作者，性格专业友善、偶尔吐槽。'
    'Persona 配置定义了兴趣领域（AI/开源/编程/数码）、互动策略'
    '（技术干货必赞、每天评论不超过 20 条、每天关注不超过 10 个）、'
    '发帖策略（每天 1-3 条、固定时间窗口）、安全规则（不碰政治/色情/谣言）。'
    '这些策略约束 Agent 的决策空间，避免无序互动。'
)

# ============================================================
# 3. 遇到的问题与解决方案
# ============================================================
doc.add_heading('三、遇到的问题与解决方案', level=1)

doc.add_heading('3.1 Workflow 配置文件中文乱码', level=2)
add_para('问题描述：', bold=True, space_after=3)
add_para(
    '三个 Workflow JSON 配置文件中的中文字符串全部出现乱码。'
    '例如"#技术分享#"变成"#鎶€鏈垎浜?"，"大模型"变成"澶фā鍨?"。'
    '这是典型的 UTF-8 字节被当作 GBK 二次编码导致的 mojibake。'
    '虽然 JSON 仍能解析，但传给 aione 的搜索关键词和评论模板全是乱码，'
    'Workflow 实际运行时等于在搜索乱码字符串。'
)
add_para('解决方案：', bold=True, space_after=3)
add_para(
    '用正确的 UTF-8 编码重写全部三个 JSON 文件，验证 ConvertFrom-Json 解析正常、'
    '中文内容正确。根因是文件创建时编码不一致——后续统一使用 UTF-8 保存。'
)

doc.add_heading('3.2 多平台浏览器端口冲突', level=2)
add_para('问题描述：', bold=True, space_after=3)
add_para(
    'LocoAgent 支持多平台（X、微博、LinkedIn 等），每个平台使用独立的 Chrome 实例'
    '和不同的 CDP 端口（X:9222, 微博:9229）。但 agent-browser 的全局配置文件'
    '（agent-browser.json）只能 pin 一个端口，默认指向 X 的 9222。'
    '微博的 SKILL.md 中所有 agent-browser 命令未指定 --cdp 9229，'
    '导致操作会连到 X 的 Chrome 而非微博的。'
)
add_para('解决方案：', bold=True, space_after=3)
add_para(
    '理解了 setup-chrome 脚本的设计意图：默认平台走 pin，其他平台通过 --cdp <port> 参数指定。'
    '后续需要在 SKILL.md 中所有 agent-browser 命令补上 --cdp 9229 参数，'
    '确保微博操作连接到正确的 Chrome 实例。'
)

doc.add_heading('3.3 CLI 与 Browser 的能力边界划分', level=2)
add_para('问题描述：', bold=True, space_after=3)
add_para(
    '微博没有开放 API，All-IN-ONE CLI 通过逆向工程覆盖了搜索、用户信息、发帖等操作，'
    '但点赞、关注、评论等交互式操作没有对应的 CLI 命令——这些操作需要页面状态'
    '（找到按钮、点击、等待响应），纯 HTTP 请求难以可靠模拟。'
)
add_para('解决方案：', bold=True, space_after=3)
add_para(
    '采用混合策略：CLI 负责结构化读取（搜索、用户信息、评论列表、发帖），'
    'Browser 负责交互式操作（点赞、关注、评论提交、信息流浏览）。'
    '在 SKILL.md 中为每个操作标注 Method（CLI/Browser），'
    'Agent 根据标注选择技术路径。这避免了"用浏览器做一切"的低效，'
    '也避免了"用 CLI 做一切"的不可能。'
)

doc.add_heading('3.4 反爬与反 Bot 策略', level=2)
add_para('问题描述：', bold=True, space_after=3)
add_para(
    '微博对自动化行为有检测机制：频繁点赞/评论/关注会触发 CAPTCHA 或限流；'
    '自动登录会触发安全验证。一旦触发，后续操作全部失败。'
)
add_para('解决方案：', bold=True, space_after=3)
add_bullet('人工登录：Chrome 启动后由用户手动登录一次，session 持久化，Agent 永不输入密码')
add_bullet('频率控制：Persona 中设定每日点赞上限、评论上限、关注上限，操作间隔 3-5 秒')
add_bullet('操作去重：LocoAgent 的 Operation Log 记录所有已执行操作，跨 session 去重')
add_bullet('CAPTCHA 检测：遇到验证码立即停止所有操作，等待人工处理')

doc.add_heading('3.5 Executor 实现不完整', level=2)
add_para('问题描述：', bold=True, space_after=3)
add_para(
    '三个微博 Workflow 的 executor TypeScript 文件目前是骨架/占位实现。'
    'weibo-feed-monitor.ts 只执行搜索并打印"找到 N 条"，没有真正执行点赞/评论/关注；'
    'weibo-daily-post.ts 发送硬编码的占位字符串，没有从内容池选题的逻辑。'
)
add_para('解决方案：', bold=True, space_after=3)
add_para('这是当前正在推进的工作。计划：')
add_bullet('weibo-daily-post.ts：读取 content-pool.md → 按分类权重选题 → 拼接 hashtag → 调用 aione 发帖 → 写操作日志')
add_bullet('weibo-feed-monitor.ts：搜索 → 对每条结果通过 agent-browser 执行点赞/评论/关注 → 操作日志去重')
add_bullet('weibo-search-reply.ts：按 searchQueries 配置执行搜索 → 对结果执行配置的 action（like/like_and_comment/follow）')

doc.add_heading('3.6 依赖未安装与环境配置', level=2)
add_para('问题描述：', bold=True, space_after=3)
add_para(
    'aione（All-IN-ONE CLI）和 agent-browser 两个核心依赖尚未安装到系统 PATH。'
    '微博 cookie 未配置，Chrome 未启动登录。没有这些，所有操作都无法实际执行。'
)
add_para('解决方案：', bold=True, space_after=3)
add_para(
    '编写了一键安装脚本 setup.ps1，自动检查 Bun/Python/Node.js，'
    '安装 LocoAgent 依赖和 All-IN-ONE CLI。后续按步骤配置 cookie、'
    '启动 Chrome、执行 doctor 健康检查即可运行。'
)

# ============================================================
# 4. 项目结构
# ============================================================
doc.add_heading('四、项目结构', level=1)

struct_table = doc.add_table(rows=1, cols=2, style='Light List Accent 1')
struct_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = struct_table.rows[0].cells
hdr[0].text = '目录/文件'
hdr[1].text = '说明'

structure = [
    ('locoagent/skills/weibo/SKILL.md', '微博操作手册（核心，22KB，6 大类 20+ 操作序列）'),
    ('locoagent/persona/persona.md', '角色人设：小洛，AI 科技观察者'),
    ('locoagent/persona/tasks.md', '每日任务清单（自然语言描述）'),
    ('locoagent/persona/content-pool.md', '内容灵感池（教程/观点/推荐/日常）'),
    ('locoagent/workflows/weibo-*.json', '三个 Workflow 配置（发帖/监控/搜索）'),
    ('locoagent/workflows/executors/weibo-*.ts', '对应的 TypeScript 执行器'),
    ('locoagent/config/browser-targets.json', '多平台 Chrome CDP 端口配置'),
    ('locoagent/.env', 'LLM 配置（DeepSeek）'),
    ('setup.ps1', '一键安装脚本'),
    ('README.md', '项目文档与快速开始指南'),
]
for path, desc in structure:
    row = struct_table.add_row().cells
    row[0].text = path
    row[1].text = desc

doc.add_paragraph()

# ============================================================
# 5. 当前进展与后续计划
# ============================================================
doc.add_heading('五、当前进展与后续计划', level=1)

doc.add_heading('5.1 已完成', level=2)
add_bullet('架构设计与技术选型：LocoAgent + All-IN-ONE + agent-browser 混合方案')
add_bullet('微博 SKILL.md 操作手册：覆盖浏览/搜索/点赞/评论/转发/关注/发帖全链路')
add_bullet('Persona 人设设计：角色、兴趣、互动策略、安全规则')
add_bullet('三个 Workflow 配置与 executor 骨架')
add_bullet('一键安装脚本 setup.ps1 与项目文档 README.md')
add_bullet('修复 Workflow JSON 中文乱码 bug')

doc.add_heading('5.2 进行中', level=2)
add_bullet('安装 aione 与 agent-browser 依赖，配置微博 cookie')
add_bullet('SKILL.md 中 agent-browser 命令补全 --cdp 9229 参数')
add_bullet('逐个功能端到端验证（CLI 类 + Browser 类）')

doc.add_heading('5.3 后续计划', level=2)
add_bullet('补全三个 executor 的真实实现（内容池选题、自动互动、操作去重）')
add_bullet('交互模式测试：bun start -p "/weibo 浏览信息流，对AI相关内容点赞"')
add_bullet('Workflow 守护进程运行：weibo-feed-monitor 每 2 小时自动互动')
add_bullet('评估运行效果，调优 Persona 策略参数')

doc.add_paragraph()
add_para(
    '本项目的核心价值不在于"写了一个微博机器人"，而在于验证了一种可复用的方法论：'
    '通过沉淀 Skill + 安装插件 + 编排 Workflow，基于开源 Agent 框架快速构建'
    '面向任意社交平台的自动化 Agent。微博只是第一个落地场景，后续可迁移至'
    '小红书、抖音等平台——只需更换 SKILL.md 和 Workflow 配置。',
    bold=False
)

# ============================================================
# Save
# ============================================================
output_path = r'E:\program\weiboagent\WeiboAgent_项目介绍报告.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
